
import json
from copy import deepcopy
from datetime import datetime
from io import BytesIO

import streamlit as st
from docx import Document


def _safe_str(v, default=""):
    if v is None:
        return default
    if isinstance(v, str):
        return v
    try:
        return str(v)
    except Exception:
        return default


def _safe_list(v):
    if v is None:
        return []
    if isinstance(v, list):
        return v
    if isinstance(v, tuple):
        return list(v)
    return [v]


def _normalize_text_list(v):
    out = []
    for x in _safe_list(v):
        if isinstance(x, dict):
            title = _safe_str(x.get("title", "")).strip()
            discussion = _safe_str(x.get("discussion", x.get("summary", x.get("text", "")))).strip()
            merged = ": ".join([p for p in [title, discussion] if p]).strip()
            if merged:
                out.append(merged)
        else:
            s = _safe_str(x).strip()
            if s:
                out.append(s)
    return out


def _normalize_action_items(v):
    items = []
    for item in _safe_list(v):
        if isinstance(item, dict):
            items.append({
                "task": _safe_str(item.get("task", "")).strip(),
                "assignee": _safe_str(item.get("assignee", "")).strip() or None,
                "deadline": _safe_str(item.get("deadline", "")).strip() or None,
                "details": _safe_str(item.get("details", "")).strip(),
            })
        else:
            txt = _safe_str(item).strip()
            if txt:
                items.append({"task": txt, "assignee": None, "deadline": None, "details": ""})
    return items


def _normalize_agenda_items(v):
    items = []
    for item in _safe_list(v):
        if isinstance(item, dict):
            items.append({
                "item_no": _safe_str(item.get("item_no", "")).strip(),
                "title": _safe_str(item.get("title", "")).strip(),
                "discussion": _safe_str(item.get("discussion", "")).strip(),
                "decisions": _normalize_text_list(item.get("decisions", [])),
                "action_items": _normalize_action_items(item.get("action_items", [])),
            })
        else:
            txt = _safe_str(item).strip()
            if txt:
                items.append({
                    "item_no": "",
                    "title": txt,
                    "discussion": "",
                    "decisions": [],
                    "action_items": [],
                })
    return items


def _sanitize_minutes(m):
    m = deepcopy(m) if isinstance(m, dict) else {}
    m.setdefault("summary", "")
    m["summary"] = _safe_str(m.get("summary", "")).strip()
    m["topics"] = _normalize_text_list(m.get("topics", []))
    m["decisions"] = _normalize_text_list(m.get("decisions", []))
    m["key_issues"] = _normalize_text_list(m.get("key_issues", []))
    m["action_items"] = _normalize_action_items(m.get("action_items", []))
    m["agenda_items"] = _normalize_agenda_items(m.get("agenda_items", []))
    m["other_matters"] = _safe_str(m.get("other_matters", "")).strip()
    nm = m.get("next_meeting", None)
    m["next_meeting"] = None if nm in (None, "", "null") else _safe_str(nm).strip()
    return m


def _minutes_to_text(m, meeting_info=None):
    m = _sanitize_minutes(m)
    meeting_info = meeting_info or {}
    lines = []
    title = meeting_info.get("meeting_name") or "會議紀錄"
    lines.append(title)
    date = _safe_str(meeting_info.get("date", "")).strip()
    venue = _safe_str(meeting_info.get("venue", "")).strip()
    attendees = _safe_str(meeting_info.get("attendees", "")).strip()
    if date:
        lines.append(f"日期：{date}")
    if venue:
        lines.append(f"地點：{venue}")
    if attendees:
        lines.append(f"出席：{attendees}")
    lines.append("")
    lines.append("【執行摘要】")
    lines.append(m.get("summary") or "—")
    lines.append("")
    if m.get("agenda_items"):
        lines.append("【議程討論紀錄】")
        for item in m["agenda_items"]:
            head = " ".join([x for x in [item.get("item_no", "").strip(), item.get("title", "").strip()] if x]).strip()
            if head:
                lines.append(head)
            lines.append(item.get("discussion") or "（會議中未有討論）")
            for d in item.get("decisions", []):
                lines.append(f"- 決議：{d}")
            for a in item.get("action_items", []):
                task = a.get("task") or "—"
                assignee = a.get("assignee") or "待定"
                deadline = a.get("deadline") or "待定"
                details = a.get("details") or ""
                line = f"- 跟進：{task}（負責：{assignee}；期限：{deadline}）"
                if details:
                    line += f"；說明：{details}"
                lines.append(line)
            lines.append("")
        if m.get("other_matters"):
            lines.append("【其他事項】")
            lines.append(m["other_matters"])
            lines.append("")
    else:
        if m.get("topics"):
            lines.append("【主要議題】")
            for t in m["topics"]:
                lines.append(f"- {t}")
            lines.append("")
        if m.get("decisions"):
            lines.append("【決議事項】")
            for d in m["decisions"]:
                lines.append(f"- {d}")
            lines.append("")
        if m.get("action_items"):
            lines.append("【行動項目】")
            for i, a in enumerate(m["action_items"], 1):
                task = a.get("task") or "—"
                assignee = a.get("assignee") or "待定"
                deadline = a.get("deadline") or "待定"
                details = a.get("details") or ""
                line = f"{i}. {task}（負責：{assignee}；期限：{deadline}）"
                if details:
                    line += f"；說明：{details}"
                lines.append(line)
            lines.append("")
    if m.get("key_issues"):
        lines.append("【待跟進事項】")
        for k in m["key_issues"]:
            lines.append(f"- {k}")
        lines.append("")
    if m.get("next_meeting"):
        lines.append("【下次會議】")
        lines.append(m["next_meeting"])
        lines.append("")
    return "\n".join(lines).strip()


def _build_docx_bytes(m, meeting_info=None):
    m = _sanitize_minutes(m)
    meeting_info = meeting_info or {}
    doc = Document()
    doc.add_heading(meeting_info.get("meeting_name") or "會議紀錄", level=0)
    meta = []
    for label, key in [("日期", "date"), ("地點", "venue"), ("出席", "attendees")]:
        val = _safe_str(meeting_info.get(key, "")).strip()
        if val:
            meta.append(f"{label}：{val}")
    if meta:
        doc.add_paragraph("\n".join(meta))
    doc.add_heading("執行摘要", level=1)
    doc.add_paragraph(m.get("summary") or "—")
    if m.get("agenda_items"):
        doc.add_heading("議程討論紀錄", level=1)
        for item in m["agenda_items"]:
            head = " ".join([x for x in [item.get("item_no", "").strip(), item.get("title", "").strip()] if x]).strip() or "未命名議項"
            doc.add_heading(head, level=2)
            doc.add_paragraph(item.get("discussion") or "（會議中未有討論）")
            if item.get("decisions"):
                doc.add_paragraph("決議事項")
                for d in item["decisions"]:
                    doc.add_paragraph(d, style="List Bullet")
            if item.get("action_items"):
                doc.add_paragraph("行動項目")
                for a in item["action_items"]:
                    task = a.get("task") or "—"
                    assignee = a.get("assignee") or "待定"
                    deadline = a.get("deadline") or "待定"
                    details = a.get("details") or ""
                    txt = f"{task}（負責：{assignee}；期限：{deadline}）"
                    if details:
                        txt += f"；說明：{details}"
                    doc.add_paragraph(txt, style="List Bullet")
        if m.get("other_matters"):
            doc.add_heading("其他事項", level=1)
            doc.add_paragraph(m["other_matters"])
    else:
        if m.get("topics"):
            doc.add_heading("主要議題", level=1)
            for t in m["topics"]:
                doc.add_paragraph(t, style="List Bullet")
        if m.get("decisions"):
            doc.add_heading("決議事項", level=1)
            for d in m["decisions"]:
                doc.add_paragraph(d, style="List Bullet")
        if m.get("action_items"):
            doc.add_heading("行動項目", level=1)
            for a in m["action_items"]:
                task = a.get("task") or "—"
                assignee = a.get("assignee") or "待定"
                deadline = a.get("deadline") or "待定"
                details = a.get("details") or ""
                txt = f"{task}（負責：{assignee}；期限：{deadline}）"
                if details:
                    txt += f"；說明：{details}"
                doc.add_paragraph(txt, style="List Bullet")
    if m.get("key_issues"):
        doc.add_heading("待跟進事項", level=1)
        for k in m["key_issues"]:
            doc.add_paragraph(k, style="List Bullet")
    if m.get("next_meeting"):
        doc.add_heading("下次會議", level=1)
        doc.add_paragraph(m["next_meeting"])
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def render_minutes(minutes: dict, meeting_info: dict | None = None, key_prefix: str = "minutes"):
    meeting_info = meeting_info or {}
    safe = _sanitize_minutes(minutes)
    box = st.container()
    with box:
        st.subheader("📝 會議紀錄")
        if not safe:
            st.info("尚未有會議紀錄內容。")
            return None

        has_agenda = bool(safe.get("agenda_items"))
        if has_agenda:
            st.caption(f"議程模式：共 {len(safe['agenda_items'])} 個議項")

        summary_key = f"{key_prefix}_summary"
        topics_key = f"{key_prefix}_topics"
        decisions_key = f"{key_prefix}_decisions"
        key_issues_key = f"{key_prefix}_key_issues"
        next_meeting_key = f"{key_prefix}_next_meeting"
        other_matters_key = f"{key_prefix}_other_matters"

        safe["summary"] = st.text_area("執行摘要", value=safe.get("summary", ""), height=120, key=summary_key)

        if has_agenda:
            edited_agenda = []
            for idx, item in enumerate(safe["agenda_items"]):
                with st.expander(f"議項 {idx + 1}：{' '.join([x for x in [item.get('item_no',''), item.get('title','')] if x]).strip() or '未命名議項'}", expanded=(idx == 0)):
                    item_no = st.text_input("項次", value=item.get("item_no", ""), key=f"{key_prefix}_ag_no_{idx}")
                    title = st.text_input("標題", value=item.get("title", ""), key=f"{key_prefix}_ag_title_{idx}")
                    discussion = st.text_area("討論內容", value=item.get("discussion", ""), height=140, key=f"{key_prefix}_ag_disc_{idx}")
                    decisions_text = st.text_area("決議事項（每行一項）", value="\n".join(item.get("decisions", [])), height=100, key=f"{key_prefix}_ag_dec_{idx}")
                    actions_text = st.text_area("行動項目 JSON（可留空）", value=json.dumps(item.get("action_items", []), ensure_ascii=False, indent=2), height=160, key=f"{key_prefix}_ag_act_{idx}")
                    try:
                        action_items = _normalize_action_items(json.loads(actions_text)) if actions_text.strip() else []
                    except Exception:
                        st.warning("行動項目 JSON 格式無效，已保留原值。")
                        action_items = item.get("action_items", [])
                    edited_agenda.append({
                        "item_no": item_no,
                        "title": title,
                        "discussion": discussion,
                        "decisions": [x.strip() for x in decisions_text.splitlines() if x.strip()],
                        "action_items": action_items,
                    })
            safe["agenda_items"] = edited_agenda
            safe["other_matters"] = st.text_area("其他事項", value=safe.get("other_matters", ""), height=100, key=other_matters_key)
            flat_decisions = []
            flat_actions = []
            flat_topics = []
            for item in safe["agenda_items"]:
                head = " ".join([x for x in [item.get("item_no", ""), item.get("title", "")] if x]).strip()
                flat_topics.append(f"{head}: {item.get('discussion') or '（會議中未有討論）'}".strip())
                flat_decisions.extend(item.get("decisions", []))
                flat_actions.extend(item.get("action_items", []))
            safe["topics"] = flat_topics
            safe["decisions"] = flat_decisions
            safe["action_items"] = flat_actions
        else:
            safe["topics"] = [x.strip() for x in st.text_area("主要議題（每行一項）", value="\n".join(safe.get("topics", [])), height=140, key=topics_key).splitlines() if x.strip()]
            safe["decisions"] = [x.strip() for x in st.text_area("決議事項（每行一項）", value="\n".join(safe.get("decisions", [])), height=140, key=decisions_key).splitlines() if x.strip()]
            action_text = st.text_area("行動項目 JSON", value=json.dumps(safe.get("action_items", []), ensure_ascii=False, indent=2), height=180, key=f"{key_prefix}_actions")
            try:
                safe["action_items"] = _normalize_action_items(json.loads(action_text)) if action_text.strip() else []
            except Exception:
                st.warning("行動項目 JSON 格式無效，已保留原值。")

        safe["key_issues"] = [x.strip() for x in st.text_area("待跟進事項（每行一項）", value="\n".join(safe.get("key_issues", [])), height=100, key=key_issues_key).splitlines() if x.strip()]
        nm_val = "" if safe.get("next_meeting") is None else _safe_str(safe.get("next_meeting"))
        nm_out = st.text_input("下次會議", value=nm_val, key=next_meeting_key)
        safe["next_meeting"] = nm_out.strip() or None

        text_out = _minutes_to_text(safe, meeting_info)
        st.text_area("預覽文字", value=text_out, height=320, key=f"{key_prefix}_preview")

        c1, c2, c3 = st.columns(3)
        with c1:
            st.download_button("⬇️ 下載 TXT", data=text_out.encode("utf-8-sig"), file_name=f"minutes_{key_prefix}.txt", mime="text/plain", key=f"{key_prefix}_dl_txt")
        with c2:
            st.download_button("⬇️ 下載 JSON", data=json.dumps(safe, ensure_ascii=False, indent=2).encode("utf-8-sig"), file_name=f"minutes_{key_prefix}.json", mime="application/json", key=f"{key_prefix}_dl_json")
        with c3:
            st.download_button("⬇️ 下載 DOCX", data=_build_docx_bytes(safe, meeting_info), file_name=f"minutes_{key_prefix}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"{key_prefix}_dl_docx")

        return safe
