"""
Export Service — DOCX, TXT, SRT
改進：
  1. 字體後備鏈：按平台自動選擇 CJK 字體（Windows/macOS/Linux）
  2. 自動設定東亞字體（避免 macOS 出現方框）
  3. 頁尾加入頁碼及生成時間戳
"""
import io, sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _get_cjk_font() -> str:
    """按作業系統回傳最合適的 CJK 字體。"""
    if sys.platform == "darwin":
        return "PingFang HK"
    elif sys.platform == "win32":
        return "Microsoft JhengHei"
    return "Noto Sans CJK TC"


def _set_doc_font(doc: Document, font_name: str):
    """設定 Normal 樣式的中西文字體。"""
    style    = doc.styles["Normal"]
    style.font.name  = font_name
    style.font.size  = Pt(11)
    rPr    = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:hAnsi"),    font_name)


def _add_footer(doc: Document):
    """頁尾：生成時間 + 「第 X 頁」。"""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.clear()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ts_run = para.add_run(f"生成時間：{datetime.now().strftime('%Y-%m-%d %H:%M')}　第 ")
    ts_run.font.size = Pt(8)

    # 插入 PAGE 欄位
    fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
    instr     = OxmlElement("w:instrText"); instr.text = "PAGE"
    fld_end   = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
    pg_run = para.add_run(" 頁")
    pg_run._r.insert(0, fld_end)
    pg_run._r.insert(0, instr)
    pg_run._r.insert(0, fld_begin)
    pg_run.font.size = Pt(8)


def _s(v, fallback="—") -> str:
    if v is None or str(v).strip() == "":
        return fallback
    return str(v).strip()


def export_minutes_docx(
    minutes: dict,
    meeting_info: dict = None,
    version_note: str = "",
) -> bytes:
    doc       = Document()
    font_name = _get_cjk_font()
    _set_doc_font(doc, font_name)
    _add_footer(doc)

    mi    = meeting_info or {}
    title = mi.get("meeting_name", "會議紀錄")
    h     = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 會議資訊
    for label, key in [("日期", "date"), ("地點", "venue"),
                        ("出席", "attendees"), ("會議類型", "meeting_type")]:
        val = _s(mi.get(key))
        if val != "—":
            p = doc.add_paragraph()
            run = p.add_run(f"{label}："); run.bold = True
            p.add_run(val)

    if version_note:
        p   = doc.add_paragraph()
        run = p.add_run(f"版本：{version_note}")
        run.italic = True
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # 執行摘要
    doc.add_heading("執行摘要", level=2)
    doc.add_paragraph(_s(minutes.get("summary")))

    # 議程模式
    if minutes.get("_has_agenda") and minutes.get("agenda_items"):
        doc.add_heading("議程討論紀錄", level=2)
        for item in minutes.get("agenda_items", []):
            no   = _s(item.get("item_no"), "")
            ttl  = _s(item.get("title"),   "（無標題）")
            disc = _s(item.get("discussion"), "（未有討論）")
            doc.add_heading(f"{no} {ttl}".strip(), level=3)
            doc.add_paragraph(disc)
            for d in item.get("decisions", []):
                if _s(d) != "—":
                    doc.add_paragraph(f"✓ {d}", style="List Bullet")
            for a in item.get("action_items", []):
                task = _s(a.get("task"))
                who  = _s(a.get("assignee"), "待定")
                dl   = _s(a.get("deadline"),  "待定")
                det  = _s(a.get("details"),   "")
                txt  = f"• {task}（{who} / {dl}）"
                if det and det != "—":
                    txt += f" — {det}"
                doc.add_paragraph(txt)
        om = _s(minutes.get("other_matters"), "")
        if om and om != "—":
            doc.add_heading("其他事項", level=2)
            doc.add_paragraph(om)
    else:
        # 標準模式
        for section_title, key in [("主要議題", "topics"), ("決議事項", "decisions")]:
            items = [i for i in minutes.get(key, []) if _s(i) != "—"]
            if items:
                doc.add_heading(section_title, level=2)
                for item in items:
                    doc.add_paragraph(f"• {item}", style="List Bullet")

        ai = [a for a in minutes.get("action_items", []) if _s(a.get("task")) != "—"]
        if ai:
            doc.add_heading("行動項目", level=2)
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h_txt in enumerate(["#", "任務", "負責人", "截止日期", "備注"]):
                hdr[i].text = h_txt
                hdr[i].paragraphs[0].runs[0].bold = True
            for idx, item in enumerate(ai, 1):
                row = table.add_row().cells
                row[0].text = str(idx)
                row[1].text = _s(item.get("task"))
                row[2].text = _s(item.get("assignee"), "待定")
                row[3].text = _s(item.get("deadline"),  "待定")
                row[4].text = _s(item.get("details"),   "")

    # 待跟進
    ki = [q for q in minutes.get("key_issues", []) if _s(q) != "—"]
    if ki:
        doc.add_heading("待跟進事項", level=2)
        for q in ki:
            doc.add_paragraph(f"⚠ {q}", style="List Bullet")

    nm = _s(minutes.get("next_meeting"), "")
    if nm and nm != "—":
        doc.add_heading("下次會議", level=2)
        doc.add_paragraph(nm)

    # PTA 問答（家長教師會專用）
    pta = minutes.get("pta_qna", [])
    if pta:
        doc.add_heading("家長問答紀錄", level=2)
        for qa in pta:
            asker  = _s(qa.get("asker"), "家長")
            q_text = _s(qa.get("question"), "")
            r_text = _s(qa.get("response"), "")
            p = doc.add_paragraph()
            p.add_run(f"【{asker}】").bold = True
            p.add_run(q_text)
            p2 = doc.add_paragraph()
            p2.add_run("　校方回應：").bold = True
            p2.add_run(r_text)

    gen = _s(minutes.get("_generated_by"), "")
    if gen and gen != "—":
        doc.add_paragraph()
        p = doc.add_paragraph(f"（由 {gen} 自動生成）")
        p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.runs[0].font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_transcript_docx(segments: list, full_text: str) -> bytes:
    doc       = Document()
    font_name = _get_cjk_font()
    _set_doc_font(doc, font_name)
    _add_footer(doc)
    doc.add_heading("逐字稿", level=1)
    for seg in segments:
        m1, s1 = divmod(int(seg.get("start", 0)), 60)
        m2, s2 = divmod(int(seg.get("end",   0)), 60)
        spk    = f"[{seg['speaker']}] " if seg.get("speaker") else ""
        p      = doc.add_paragraph()
        p.add_run(f"[{m1:02d}:{s1:02d}–{m2:02d}:{s2:02d}] {spk}").bold = True
        p.add_run(_s(seg.get("text"), ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_srt(segments: list) -> str:
    def fmt(s):
        h  = int(s // 3600)
        m  = int((s % 3600) // 60)
        sc = int(s % 60)
        ms = int((s % 1) * 1000)
        return f"{h:02d}:{m:02d}:{sc:02d},{ms:03d}"
    lines = []
    for i, seg in enumerate(segments, 1):
        spk = f"[{seg['speaker']}] " if seg.get("speaker") else ""
        lines += [
            str(i),
            f"{fmt(seg['start'])} --> {fmt(seg['end'])}",
            f"{spk}{seg.get('text', '')}",
            "",
        ]
    return "\n".join(lines)
